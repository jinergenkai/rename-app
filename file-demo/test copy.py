import os
import re
from docx import Document

# Äá»c danh sÃ¡ch tá»« khÃ³a prefix tá»« file match.txt
def load_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        keywords = [line.strip() for line in f.readlines()]
    return keywords

# Äá»c danh sÃ¡ch tá»« khÃ³a cáº§n bá» qua tá»« file ignore.txt
def load_ignore_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        ignore_keywords = [line.strip() for line in f.readlines()]
    return ignore_keywords

# ÄÃ¡nh Ä‘iá»ƒm cÃ¡c dÃ²ng theo cÃ¡c tiÃªu chÃ­
def score_line(line, match_keywords, ignore_keywords, font_sizes):
    score = 0
    # 1. Náº¿u dÃ²ng cÃ³ tá»« khÃ³a trÃ¹ng vá»›i match.txt, cá»™ng 1 Ä‘iá»ƒm
    if any(keyword.lower() in line.lower() for keyword in match_keywords):
        score += 1
    
    # 2. Náº¿u dÃ²ng lÃ  cÄƒn giá»¯a, cá»™ng 1 Ä‘iá»ƒm
    if line == line.strip() and line.strip() == line.strip().center(len(line)):
        score += 1
    
    # 3. Náº¿u dÃ²ng viáº¿t hoa háº¿t, cá»™ng 1 Ä‘iá»ƒm
    if line.isupper():
        score += 1
    
    # 4. Náº¿u dÃ²ng cÃ³ font chá»¯ lá»›n nháº¥t trong 20 dÃ²ng Ä‘áº§u tiÃªn, cá»™ng 1 Ä‘iá»ƒm
    if font_sizes and font_sizes.get(line, 0) == max(font_sizes.values()):
        score += 1
    
    # 5. Náº¿u dÃ²ng chá»©a tá»« khÃ³a trong ignore.txt, trá»« 1 Ä‘iá»ƒm
    if any(ignore.lower() in line.lower() for ignore in ignore_keywords):
        score -= 1
    
    return score

# Xá»­ lÃ½ vÄƒn báº£n vÃ  tÃ¬m tiÃªu Ä‘á»
def process_text(doc, text, match_keywords, ignore_keywords):
    lines = text.split("\n")
    
    # LÆ°u trá»¯ cÃ¡c thuá»™c tÃ­nh cá»§a cÃ¡c dÃ²ng
    font_sizes = {}
    for para in doc.paragraphs[:20]:
        for run in para.runs:
            line = para.text.strip()
            if line:  # Náº¿u Ä‘oáº¡n khÃ´ng rá»—ng
                font_sizes[line] = max(font_sizes.get(line, 0), run.font.size.pt if run.font.size else 0)
    
    # ÄÃ¡nh Ä‘iá»ƒm cho tá»«ng dÃ²ng
    scored_lines = []
    for line in lines[:20]:
        score = score_line(line, match_keywords, ignore_keywords, font_sizes)
        scored_lines.append((line, score))
    
    # TÃ¬m dÃ²ng cÃ³ Ä‘iá»ƒm cao nháº¥t lÃ m tiÃªu Ä‘á»
    best_title = max(scored_lines, key=lambda x: x[1])[0]
    
    return best_title.strip()

# Äá»c file DOCX
def read_docx_text(filepath):
    doc = Document(filepath)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return doc, full_text

# HÃ m Ä‘á»•i tÃªn file
def rename_file_with_rules(filepath, match_keywords, ignore_keywords):
    doc, original_text = read_docx_text(filepath)
    new_title = process_text(doc, original_text, match_keywords, ignore_keywords)
    
    # Táº¡o tÃªn file tá»« tiÃªu Ä‘á» Ä‘Ã£ tÃ¬m Ä‘Æ°á»£c
    new_filename = new_title[:200]  # Giá»›i háº¡n 200 kÃ½ tá»±
    new_filename = re.sub(r'[<>:"/\\|?*]', '', new_filename)  # Lá»c cÃ¡c kÃ½ tá»± khÃ´ng há»£p lá»‡ trong tÃªn file
    # new_path = os.path.join(os.path.dirname(filepath), new_filename + ".docx")
    
    # os.rename(filepath, new_path)
    print(f"âœ… Äá»•i tÃªn file thÃ nh: {new_filename}.docx")

# ğŸ§ª Thá»­ demo
match_keywords = load_keywords("match.txt")  # Load tá»« khÃ³a tá»« file match.txt
ignore_keywords = load_ignore_keywords("ignore.txt")  # Load tá»« khÃ³a cáº§n bá» qua tá»« file ignore.txt
rename_file_with_rules("123.docx", match_keywords, ignore_keywords)
