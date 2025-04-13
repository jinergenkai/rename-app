import os
import re
from docx import Document
import pprint
import win32com.client  # Dùng để xử lý file .doc (Windows)
from datetime import datetime

# Logging function
def log_operation(operation_type, filepath, new_name=None, error=None):
    log_file = "rename_log.txt"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    if operation_type == "ERROR":
        log_entry = f"[{timestamp}] ERROR: {filepath} - {error}\n"
    elif operation_type == "RENAME":
        log_entry = f"[{timestamp}] RENAME: {filepath} -> {new_name}\n"
    
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(log_entry)

def validate_file(filepath):
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    if not os.path.isfile(filepath):
        raise ValueError(f"Not a file: {filepath}")
    
    if not os.access(filepath, os.R_OK):
        raise PermissionError(f"No read permission: {filepath}")
    
    file_size = os.path.getsize(filepath)
    if file_size == 0:
        raise ValueError(f"File is empty: {filepath}")
    
    # 100MB limit
    if file_size > 100 * 1024 * 1024:
        raise ValueError(f"File too large (>100MB): {filepath}")

class TextParagraph:
    def __init__(self, text, font_size=None, is_centered=False):
        self.text = text.strip()
        self.points = 0
        self.is_uppercase = text.isupper()
        self.is_centered = is_centered
        self.font_size = font_size
    
    def to_dict(self):
        return {
            'text': self.text,
            'points': self.points,  # This will now reflect the current points
            'is_uppercase': self.is_uppercase,
            'is_centered': self.is_centered,
            'font_size': self.font_size
        }
    
    def add_points(self, points):
        self.points += points
    
    def matches_keyword(self, keyword):
        return keyword.lower() in self.text.lower()
    
    def has_max_font_size(self, max_font_size):
        return self.font_size == max_font_size

# Đọc danh sách từ khóa prefix từ file match.txt và ignore.txt
def load_keywords(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        keywords = [line.strip() for line in f.readlines()]
    return keywords

# Lấy font size lớn nhất từ dòng (paragraph)
def get_max_font_size(paragraph):
    max_font_size = 0
    for run in paragraph.runs:
        if run.font.size:  # Nếu có thông tin về kích thước font
            max_font_size = max(max_font_size, run.font.size.pt)  # Lấy giá trị lớn nhất (pt)
    return max_font_size

def read_docx_paragraphs(filepath, line_limit=None):
    doc = Document(filepath)
    paragraphs = []
    max_font_size_overall = 0
    processed_count = 0

    # Calculate max font size first from both paragraphs and tables
    for paragraph in doc.paragraphs:
        max_font_size_overall = max(max_font_size_overall, get_max_font_size(paragraph))
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:  # Only process non-empty paragraphs
            # Check if text is centered by comparing with centered version
            is_centered = text == text.center(len(text))
            font_size = get_max_font_size(paragraph)
            
            para = TextParagraph(
                text=text,
                font_size=font_size,
                is_centered=is_centered
            )
            paragraphs.append(para)
            
            processed_count += 1
            if line_limit and processed_count >= line_limit:
                break
    
    processed_count = 0
    # Process tables if we haven't hit the line limit
    if not line_limit or processed_count < line_limit:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        # For table cells, we consider them centered if they're in the middle columns
                        col_index = row._tr.xpath('./w:tc').index(cell._tc)
                        is_centered = col_index > 0 and col_index < len(row.cells) - 1
                        
                        # Get font size from the first paragraph in the cell
                        font_size = get_max_font_size(cell.paragraphs[0]) if cell.paragraphs else 0
                        
                        para = TextParagraph(
                            text=text,
                            font_size=font_size,
                            is_centered=is_centered
                        )
                        paragraphs.append(para)
                        
                        processed_count += 2
                        if line_limit and processed_count >= line_limit:
                            break
                if line_limit and processed_count >= line_limit:
                    break
            if line_limit and processed_count >= line_limit:
                break

    return paragraphs

def read_doc_paragraphs(filepath, line_limit=None):
    try:
        # Cố gắng tạo đối tượng Word nếu có Office
        abs_path = os.path.abspath(filepath)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_path)
        
        paragraphs = []
        processed_count = 0
        
        # Đọc từng đoạn văn bản
        for p in doc.Paragraphs:
            text = p.Range.Text.strip()
            if text:  # Chỉ lấy các đoạn không rỗng
                # Kiểm tra căn chỉnh đoạn văn (Alignment)
                is_centered = (p.Alignment == 1)  # 1 tương ứng với căn giữa (wdAlignParagraphCenter)
                
                para = TextParagraph(
                    text=text,
                    font_size=p.Range.Font.Size,
                    is_centered=is_centered
                )
                paragraphs.append(para)
                
                processed_count += 1
                if line_limit and processed_count >= line_limit:
                    break
        
        doc.Close()
        word.Quit()
        
        return paragraphs
    
    except ImportError:
        print("❌ Không thể import win32com.client. Đảm bảo rằng Microsoft Office đã được cài đặt.")
    
    except Exception as e:
        print(f"❌ Lỗi khi mở file .doc: {e}")

def clean_filename(filename):
    filename = re.sub(r'[\x00-\x1F]', '', filename)
    return re.sub(r'[<>:"/\\|?*\n\r\t\f\v]', '', filename).strip()

def process_text(paragraphs, match_keywords, ignore_keywords, **kwargs):
    # Find max font size
    max_font_size = max((p.font_size or 0) for p in paragraphs)

    # Score all paragraphs
    for para in paragraphs:
        # Check match and ignore keywords
        for keyword in match_keywords:
            if para.matches_keyword(keyword):
                para.add_points(0.5)
        for keyword in ignore_keywords:
            if para.matches_keyword(keyword):
                para.add_points(-0.5)
        
        # Check uppercase
        if para.is_uppercase:
            para.add_points(1)
        
        # Check max font size
        if para.has_max_font_size(max_font_size):
            para.add_points(1)
        
        # Check centering
        if para.is_centered:
            para.add_points(1)

    # Convert to list of dictionaries for printing AFTER scoring
    lines_data = [p.to_dict() for p in paragraphs]
    lines_data.sort(key=lambda x: x['points'], reverse=True)  # Sort by points
    pprint.pprint(lines_data) # show all lines with points

    # Sort paragraphs by points and get top 3
    sorted_paras = sorted(paragraphs, key=lambda p: p.points, reverse=True)
    top_texts = [clean_filename(p.text) for p in sorted_paras[:3]]

    print(f"Top texts: {top_texts}")  # Debugging output

    while len(top_texts) < 3:  # Ensure we have 3 items even if null
        top_texts.append("")

    cnt_empty = 0
    for i in top_texts:
        if i == "":
            cnt_empty += 1
    if cnt_empty == 3:
        return None
    
    
    # Find year if present
    all_text = "\n".join(p.text for p in paragraphs)
    year_match = re.search(r"\b(20\d{2})\b", all_text)
    year_str = f" {year_match.group(1)}" if year_match else ""

    def create_filename(main, secondary_texts):
        # Join secondary texts with separator
        secondary_part = " - " + " - ".join(t for t in secondary_texts if t)
        return f"{main}{year_str}{secondary_part}".strip()

    # First try with full secondary texts
    main_text = top_texts[0]
    secondary_texts = top_texts[1:3]
    
    # Create initial filename and clean invalid chars
    filename = create_filename(main_text, secondary_texts)
    filename = filename[:200] + " ★"

    return filename

def clean_text(text):
    # Remove control characters but keep newlines and tabs
    return re.sub(r'[^\x20-\x7E\n\t]', '', text)

def get_unique_filename(base_path):
    """Get unique filename by adding (1), (2), etc. if file exists"""
    if not os.path.exists(base_path):
        return base_path
        
    directory = os.path.dirname(base_path)
    filename = os.path.basename(base_path)
    name, ext = os.path.splitext(filename)
    
    counter = 1
    while os.path.exists(base_path):
        new_name = f"{name} ({counter}){ext}"
        base_path = os.path.join(directory, new_name)
        counter += 1
    
    return base_path

def rename_file_with_rules(filepath, match_keywords, ignore_keywords, line_limit=None, length_limit=200):
    # Validate input file
    if not os.path.exists(filepath):
        log_operation("ERROR", filepath, error="File not found")
        print(f"❌ Lỗi: File not found: {filepath}")
        return None
        
    file_extension = os.path.splitext(filepath)[1].lower()
    
    # Validate file extension
    if file_extension not in [".doc", ".docx"]:
        log_operation("ERROR", filepath, error="Unsupported file format")
        print(f"❌ Lỗi: Unsupported file format: {filepath}")
        return None
    
    try:
        # Read file content
        if file_extension == ".docx":
            paragraphs = read_docx_paragraphs(filepath, line_limit)
            print(f"✅ Đọc file .docx thành công: {filepath}")
        else:
            paragraphs = read_doc_paragraphs(filepath, line_limit)
            print(f"✅ Đọc file .doc thành công: {filepath}")
            
        if not paragraphs:
            log_operation("ERROR", filepath, error="No content found in file")
            print(f"❌ Lỗi: No content found in: {filepath}")
            return None
            
    except Exception as e:
        log_operation("ERROR", filepath, error=f"Failed to read file: {str(e)}")
        print(f"❌ Lỗi khi đọc file: {str(e)}")
        return None
    
    print(f"Đọc {len(paragraphs)} đoạn văn bản.")
    
    try:
        # Process text and generate new filename
        new_filename = process_text(paragraphs, match_keywords, ignore_keywords, length_limit=length_limit)
        if new_filename is None:
            log_operation("ERROR", filepath, error="Failed to generate filename")
            print(f"❌ Lỗi: Không thể tạo tên file mới.")
            return None
        new_path = os.path.join(os.path.dirname(filepath), new_filename + file_extension)
        
        # Get unique filename if target exists
        new_path = get_unique_filename(new_path)
        
        # Log rename operation before executing
        log_operation("RENAME", filepath, os.path.basename(new_path))
        
        try:
            # Perform rename
            # os.rename(filepath, new_path)
            print(f"✅ Đổi tên file thành: {os.path.basename(new_path)}")
            return new_path
            
        except Exception as e:
            log_operation("ERROR", filepath, error=f"Failed to rename: {str(e)}")
            print(f"❌ Lỗi khi đổi tên file: {str(e)}")
            return None
            
    except Exception as e:
        log_operation("ERROR", filepath, error=f"Failed to process: {str(e)}")
        print(f"❌ Lỗi khi xử lý file: {str(e)}")
        return None

# 🧪 Thử demo
match_keywords = load_keywords("match.txt")  # Load từ khóa từ file match.txt
ignore_keywords = load_keywords("ignore.txt")  # Load từ khóa cần loại bỏ từ ignore.txt

# Process files
pairs = []
filenames = ["file/123.docx", 'file/image.doc']
for filename in filenames:
    result = rename_file_with_rules(filename, match_keywords, ignore_keywords,
                                  line_limit=10, length_limit=200)
    if result:
        pairs.append([filename, result])

# Output results in nested list format
print(pairs)
